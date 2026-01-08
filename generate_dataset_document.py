from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import chess.pgn

# Create a new Document
doc = Document()

# Set document margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

# Add title
title = doc.add_heading('Lichess Elite Chess Dataset Documentation', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add subtitle
subtitle = doc.add_paragraph('lichess_elite_2022-02.pgn')
subtitle_format = subtitle.paragraph_format
subtitle_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.bold = True
subtitle.runs[0].font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()

# Dataset Information Section
info_heading = doc.add_heading('📊 Dataset Information', level=1)
info_heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

dataset_file = "lichess_elite_2022-02.pgn"
file_size_mb = os.path.getsize(dataset_file) / (1024 * 1024)

info_items = [
    f"📁 **Filename:** {dataset_file}",
    f"💾 **File Size:** {file_size_mb:.2f} MB",
    f"📅 **Period:** February 2022",
    f"🏆 **Type:** Elite Rated Games",
    f"⚡ **Source:** Lichess.org Database",
    f"🎯 **Purpose:** Training data for Chess AI neural network"
]

for item in info_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.left_indent = Inches(0.3)
    p.runs[0].font.size = Pt(11)

doc.add_paragraph()

# Description Section
desc_heading = doc.add_heading('📝 Description', level=1)
desc_heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

description = """The Lichess Elite dataset contains high-quality chess games from elite players on Lichess.org. 
This dataset is used to train the Chess AI neural network using supervised learning from master-level games. 
Each game in the dataset includes:

• Player ratings (typically 2000+ Elo)
• Complete move sequences in algebraic notation
• Game metadata (date, time control, result)
• Opening variations and tactical patterns
• Endgame techniques from strong players

The dataset is processed by the smart_database_parser.py module to extract training positions and convert 
them into neural network training data."""

desc_para = doc.add_paragraph(description)
desc_para.runs[0].font.size = Pt(10)
desc_para.paragraph_format.left_indent = Inches(0.3)
desc_para.paragraph_format.space_after = Pt(12)

doc.add_page_break()

# Sample Games Section
sample_heading = doc.add_heading('🎮 Sample Games from Dataset', level=1)
sample_heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

print("📖 Reading ALL games from dataset (this may take a while)...")
print("⏳ Processing large dataset - please be patient...")

try:
    with open(dataset_file, 'r', encoding='utf-8', errors='ignore') as pgn_file:
        game_count = 0
        
        while True:
            try:
                game = chess.pgn.read_game(pgn_file)
                if game is None:
                    break
                
                game_count += 1
                
                # Game heading
                game_header = doc.add_heading(f'Game {game_count}', level=2)
                game_header.runs[0].font.size = Pt(12)
                game_header.runs[0].font.color.rgb = RGBColor(0, 100, 0)
                
                # Game metadata
                event = game.headers.get("Event", "Unknown")
                date = game.headers.get("Date", "Unknown")
                white = game.headers.get("White", "Unknown")
                black = game.headers.get("Black", "Unknown")
                white_elo = game.headers.get("WhiteElo", "?")
                black_elo = game.headers.get("BlackElo", "?")
                result = game.headers.get("Result", "*")
                time_control = game.headers.get("TimeControl", "Unknown")
                opening = game.headers.get("Opening", "Unknown")
                
                metadata = [
                    f"🏆 Event: {event}",
                    f"📅 Date: {date}",
                    f"⚪ White: {white} ({white_elo})",
                    f"⚫ Black: {black} ({black_elo})",
                    f"🎯 Result: {result}",
                    f"⏱️  Time Control: {time_control}",
                    f"📖 Opening: {opening}"
                ]
                
                for meta_item in metadata:
                    p = doc.add_paragraph(meta_item)
                    p.paragraph_format.left_indent = Inches(0.3)
                    p.runs[0].font.size = Pt(9)
                    p.paragraph_format.space_before = Pt(2)
                    p.paragraph_format.space_after = Pt(2)
                
                # Moves
                moves_heading = doc.add_paragraph("📋 Moves:")
                moves_heading.runs[0].font.bold = True
                moves_heading.runs[0].font.size = Pt(10)
                moves_heading.paragraph_format.left_indent = Inches(0.3)
                moves_heading.paragraph_format.space_before = Pt(6)
                
                # Format moves in standard notation
                board = game.board()
                move_list = []
                move_number = 1
                
                for move in game.mainline_moves():
                    san_move = board.san(move)
                    if board.turn == chess.WHITE:
                        move_list.append(f"{move_number}. {san_move}")
                    else:
                        move_list[-1] += f" {san_move}"
                        move_number += 1
                    board.push(move)
                
                # Join moves with proper spacing
                moves_text = " ".join(move_list)
                
                # Add moves as code block
                moves_para = doc.add_paragraph()
                run = moves_para.add_run(moves_text)
                run.font.name = 'Courier New'
                run.font.size = Pt(8)
                moves_para.paragraph_format.left_indent = Inches(0.5)
                moves_para.paragraph_format.space_after = Pt(10)
                
                # Add separator between games
                separator = doc.add_paragraph("─" * 80)
                separator.paragraph_format.space_before = Pt(6)
                separator.paragraph_format.space_after = Pt(6)
                
                if game_count % 10 == 0:
                    print(f"  Processed {game_count} games...")
                
                # Add page break every 20 games to keep document organized
                if game_count % 20 == 0:
                    doc.add_page_break()
                    
            except Exception as e:
                print(f"  Error reading game {game_count + 1}: {e}")
                continue
        
        print(f"✅ Successfully processed {game_count} games")
        
except Exception as e:
    error_para = doc.add_paragraph(f"❌ Error reading dataset: {str(e)}")
    error_para.runs[0].font.color.rgb = RGBColor(200, 0, 0)
    print(f"❌ Error: {e}")

# Add page break before statistics
doc.add_page_break()

# Statistics Section
stats_heading = doc.add_heading('📈 Dataset Statistics', level=1)
stats_heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

stats_info = """This dataset is used by the intelligent training pipeline for:

1. **Supervised Learning Phase**
   - Extract positions from master games
   - Learn evaluation patterns
   - Train policy network on expert moves

2. **Feature Extraction**
   - Opening repertoire analysis
   - Middlegame tactical patterns
   - Endgame technique learning

3. **Quality Filtering**
   - Minimum rating threshold: 1500+
   - Game length filtering: 10-200 moves
   - Result validation (no unfinished games)
   - Time control filtering

4. **Processing Pipeline**
   - Parse PGN format
   - Convert to tensor representations
   - Generate policy targets from actual moves
   - Calculate position evaluations
   - Create training batches

The processed data feeds into the neural network training system, providing high-quality 
examples for the AI to learn chess strategy and tactics from elite players."""

stats_para = doc.add_paragraph(stats_info)
stats_para.runs[0].font.size = Pt(10)
stats_para.paragraph_format.left_indent = Inches(0.3)

# Usage Section
usage_heading = doc.add_heading('💻 Usage in Project', level=1)
usage_heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)

usage_text = """The dataset is processed by the following modules:

• **smart_database_parser.py** - Parses PGN and extracts positions
• **intelligent_training_pipeline.py** - Orchestrates the training process
• **neural_network.py** - Consumes processed training data

To use this dataset:
1. Place the PGN file in the project directory
2. Run the intelligent training pipeline
3. The parser automatically processes games
4. Extracted positions are saved for neural network training"""

usage_para = doc.add_paragraph(usage_text)
usage_para.runs[0].font.size = Pt(10)
usage_para.paragraph_format.left_indent = Inches(0.3)

# Save the document
output_filename = 'Lichess_Elite_Dataset_Documentation.docx'
doc.save(output_filename)

print("\n" + "="*60)
print("✅ Dataset documentation created successfully!")
print(f"📄 Filename: {output_filename}")
print(f"📊 Dataset: {dataset_file} ({file_size_mb:.2f} MB)")
print("="*60)
print(f"\n💾 Document saved in: {os.getcwd()}")
